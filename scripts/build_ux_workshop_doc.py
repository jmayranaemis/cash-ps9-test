from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/atelier-ux-gestionnaires-cash-alimentaire.docx")

# compact_reference_guide + named Cash brand override
BROWN = "6B5353"
DARK = "302828"
CREAM = "F5EFE8"
SAGE = "DDE5DA"
GOLD = "B99A5B"
WHITE = "FFFFFF"
LIGHT = "F6F6F6"
GRID = "D9D2CC"
MUTED = "686363"
WIDTH = 9360
INDENT = 120


def color(value):
    return RGBColor.from_string(value)


def font(run, size=10, *, bold=False, italic=False, hex_color=DARK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color(hex_color)


def shade(cell, value):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), value)


def margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, value=GRID, size=5):
    tbl_pr = table._tbl.tblPr
    root = tbl_pr.find(qn("w:tblBorders"))
    if root is None:
        root = OxmlElement("w:tblBorders")
        tbl_pr.append(root)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = root.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            root.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), value)


def geometry(table, widths, indent=INDENT):
    assert sum(widths) == WIDTH
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(WIDTH))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        no_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(no_split)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins(cell)


def cell_text(cell, text, *, size=9.2, bold=False, hex_color=DARK, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    font(p.add_run(text), size, bold=bold, hex_color=hex_color)


def table(doc, headers, rows, widths, *, size=9, header=BROWN):
    result = doc.add_table(rows=1 + len(rows), cols=len(headers))
    geometry(result, widths)
    borders(result)
    for index, title in enumerate(headers):
        shade(result.cell(0, index), header)
        cell_text(result.cell(0, index), title, size=size, bold=True, hex_color=WHITE)
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    result.rows[0]._tr.get_or_add_trPr().append(repeat)
    for row_index, values in enumerate(rows, 1):
        for col_index, value in enumerate(values):
            cell = result.cell(row_index, col_index)
            if row_index % 2 == 0:
                shade(cell, LIGHT)
            cell_text(cell, value, size=size)
    return result


def para(doc, text, *, size=10, bold=False, italic=False, hex_color=DARK,
         before=0, after=5, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    font(p.add_run(text), size, bold=bold, italic=italic, hex_color=hex_color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def callout(doc, label, text, fill=SAGE):
    box = doc.add_table(rows=1, cols=1)
    geometry(box, [WIDTH])
    borders(box, fill, 2)
    cell = box.cell(0, 0)
    shade(cell, fill)
    margins(cell, 110, 110, 160, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label + "  "), 9.5, bold=True, hex_color=BROWN)
    font(p.add_run(text), 9.5)
    para(doc, "", size=2, after=1)


def decision_line(doc, label, *, compact=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0 if compact else 3)
    p.paragraph_format.space_after = Pt(1 if compact else 4)
    font(p.add_run(label + "  "), 9.5, bold=True, hex_color=BROWN)
    font(p.add_run("______________________________________________________________"), 9.5)


def page_break(doc):
    p = para(doc, "", size=1, after=0)
    p.paragraph_format.page_break_before = True


def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.34)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = color(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, before, after in ((1, 16, 18, 10), (2, 13, 14, 7), (3, 12, 10, 5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color(BROWN if level < 3 else DARK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("CASH ALIMENTAIRE  |  UX EXPRESS"), 8, bold=True, hex_color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("Décider. Affecter. Exécuter."), 8, hex_color=MUTED)


def build():
    doc = Document()
    setup(doc)

    # Page 1 — the benchmark converted into five decisions.
    para(doc, "ATELIER UX EXPRESS", size=9, bold=True, hex_color=GOLD, after=2)
    para(doc, "90 minutes. 6 décisions. Zéro page blanche.", size=24, bold=True, hex_color=BROWN, after=5)
    para(doc, "Le benchmark est déjà fait. L’équipe valide la direction, corrige ce qui est faux et nomme les responsables.", size=11.5, after=9)

    agenda = doc.add_table(rows=1, cols=6)
    geometry(agenda, [1560] * 6, indent=0)
    borders(agenda, BROWN, 4)
    for index, (minutes, label) in enumerate([
        ("10 min", "Cible"),
        ("15 min", "Promesse"),
        ("20 min", "Accueil"),
        ("15 min", "Conversion"),
        ("15 min", "Fiche produit"),
        ("15 min", "Actions"),
    ]):
        cell = agenda.cell(0, index)
        shade(cell, CREAM if index % 2 == 0 else SAGE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(minutes), 10.5, bold=True, hex_color=BROWN)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        font(p2.add_run(label), 8.5, bold=True, hex_color=MUTED)

    heading(doc, "1. Direction issue du benchmark", 1)
    table(doc, ["Décision", "Recommandation à valider", "Verdict"], [
        ["Cible", "Métiers de bouche et professionnels de la restauration des Alpes-Maritimes.", "[ ] VALIDER\n[ ] CORRIGER"],
        ["Positionnement", "Le partenaire de terrain des métiers de bouche sur la Côte d’Azur.", "[ ] VALIDER\n[ ] CORRIGER"],
        ["CTA principal", "DEVENIR CLIENT — visible dans l’en-tête, le hero et en fin de page.", "[ ] VALIDER\n[ ] CORRIGER"],
        ["CTA secondaire", "CONSULTER LES CATALOGUES — toujours moins fort visuellement.", "[ ] VALIDER\n[ ] CORRIGER"],
        ["Différenciation", "Proximité, connaissance métier, disponibilité et service local avant le volume de références.", "[ ] VALIDER\n[ ] CORRIGER"],
    ], [1900, 5660, 1800], size=8.7)

    heading(doc, "2. Page d’accueil : ordre recommandé", 2)
    table(doc, ["#", "Bloc", "Rôle"], [
        ["1", "Hero", "Promesse + Devenir client + Catalogues"],
        ["2", "4 preuves", "Expérience, références, livraison, interlocuteur local"],
        ["3", "Familles produits", "Rendre l’offre immédiatement lisible"],
        ["4", "Entrées métiers", "Montrer que Cash comprend l’activité du client"],
        ["5", "Catalogues du moment", "Feuilleter et télécharger"],
        ["6", "Services", "Conseil, disponibilité, livraison, accompagnement"],
        ["7", "Conseil / saison", "Un contenu utile et actuel"],
        ["8", "Devenir client", "Réassurance + formulaire court"],
    ], [600, 2800, 5960], size=8.1)
    decision_line(doc, "Une seule correction indispensable :")

    # Page 2 — conversion path.
    page_break(doc)
    para(doc, "DÉCISION 4 / 6", size=9, bold=True, hex_color=GOLD, after=2)
    para(doc, "Le parcours « Devenir client »", size=22, bold=True, hex_color=BROWN, after=5)
    callout(doc, "OBJECTIF", "Une demande qualifiée en moins de 2 minutes, suivie d’un rappel commercial dans un délai réellement tenu.")

    heading(doc, "Parcours recommandé", 2)
    table(doc, ["Étape", "Ce que voit le prospect", "Décision"], [
        ["1", "Pourquoi devenir client : offre, zone, services et interlocuteur local.", "[ ] OK  [ ] À corriger"],
        ["2", "Trois étapes : demande → rappel → ouverture / accompagnement.", "[ ] OK  [ ] À corriger"],
        ["3", "Formulaire court. Aucun Kbis au premier contact sauf nécessité réelle.", "[ ] OK  [ ] À corriger"],
        ["4", "Confirmation immédiate + délai de rappel annoncé.", "[ ] OK  [ ] À corriger"],
    ], [850, 6100, 2410], size=8.8)

    heading(doc, "Champs du formulaire", 2)
    table(doc, ["Champ recommandé", "Garder", "Retirer", "Commentaire uniquement si nécessaire"], [
        ["Type d’établissement", "[ ]", "[ ]", ""],
        ["Code postal", "[ ]", "[ ]", ""],
        ["Entreprise", "[ ]", "[ ]", ""],
        ["Nom et prénom", "[ ]", "[ ]", ""],
        ["Téléphone", "[ ]", "[ ]", ""],
        ["E-mail professionnel", "[ ]", "[ ]", ""],
        ["Besoin principal", "[ ]", "[ ]", ""],
    ], [2700, 900, 900, 4860], size=8.7)

    heading(doc, "Engagement commercial — à remplir obligatoirement", 2)
    decision_line(doc, "Les demandes arrivent chez :")
    decision_line(doc, "Délai de rappel que nous pouvons tenir :")
    decision_line(doc, "En cas d’absence, relais :")
    decision_line(doc, "Message de confirmation :")
    callout(doc, "SORTIE DE CETTE PAGE", "Le formulaire est figé et une personne est responsable de chaque demande.", fill=CREAM)

    # Page 3 — product detail page, the core object of a catalogue site.
    page_break(doc)
    para(doc, "DÉCISION 5 / 6", size=9, bold=True, hex_color=GOLD, after=2)
    para(doc, "La fiche produit type", size=22, bold=True, hex_color=BROWN, after=5)
    callout(doc, "OBJECTIF", "Permettre au professionnel de vérifier en quelques secondes que le produit correspond à son usage, puis déclencher une action utile.")

    heading(doc, "Structure recommandée", 2)
    product_structure = table(doc, ["Zone", "Recommandation à valider", "Décision"], [
        ["Écran initial", "Galerie à gauche ; identité, format professionnel et CTA à droite.", "[ ] OK  [ ] Corriger"],
        ["Visuels", "Packshot net + conditionnement réel. Photo d’usage seulement si elle apporte une information.", "[ ] OK  [ ] Corriger"],
        ["Identité", "Nom, marque, référence Cash, famille, origine / labels.", "[ ] OK  [ ] Corriger"],
        ["Format pro", "Unité de vente, poids / volume, colisage et variantes visibles sans chercher.", "[ ] OK  [ ] Corriger"],
        ["Bénéfices", "3 bénéfices ou usages métier, avant la longue description.", "[ ] OK  [ ] Corriger"],
        ["CTA produit", "DEMANDER UNE INFORMATION ; la référence produit est transmise automatiquement.", "[ ] OK  [ ] Corriger"],
        ["CTA global", "DEVENIR CLIENT reste visible dans l’en-tête.", "[ ] OK  [ ] Corriger"],
        ["Sous la ligne de flottaison", "Description, préparation, conservation, composition, documents et produits associés.", "[ ] OK  [ ] Corriger"],
    ], [1800, 5600, 1960], size=7.9)
    for row in product_structure.rows:
        for cell in row.cells:
            margins(cell, 45, 45, 100, 100)

    heading(doc, "Données indispensables", 2)
    product_data = table(doc, ["Information", "Sur la page", "Dans la fiche PDF", "Indisponible"], [
        ["Conditionnement / unité de vente / colisage", "[ ]", "[ ]", "[ ]"],
        ["Origine, labels et marque", "[ ]", "[ ]", "[ ]"],
        ["Conservation, température et DLC", "[ ]", "[ ]", "[ ]"],
        ["Ingrédients, allergènes et nutrition", "[ ]", "[ ]", "[ ]"],
        ["Préparation, usage et rendement", "[ ]", "[ ]", "[ ]"],
        ["Fiche technique / fiche de sécurité", "[ ]", "[ ]", "[ ]"],
        ["Produits associés ou alternatives", "[ ]", "[ ]", "[ ]"],
    ], [5200, 1400, 1700, 1060], size=8.0)
    for row in product_data.rows:
        for cell in row.cells:
            margins(cell, 40, 40, 100, 100)

    heading(doc, "Arbitrages à fermer", 2)
    decision_line(doc, "Prix : [ ] masqué  [ ] après connexion  [ ] affiché", compact=True)
    decision_line(doc, "Disponibilité / stock affiché :", compact=True)
    decision_line(doc, "CTA produit retenu :", compact=True)
    decision_line(doc, "Source et responsable des données produit :", compact=True)
    callout(doc, "SORTIE DE CETTE PAGE", "Une fiche type est figée. Les données non disponibles sont identifiées avec un responsable.", fill=CREAM)

    # Page 4 — content and immediate execution.
    page_break(doc)
    para(doc, "DÉCISION 6 / 6", size=9, bold=True, hex_color=GOLD, after=2)
    para(doc, "Ce qu’on met en production dès maintenant", size=22, bold=True, hex_color=BROWN, after=5)
    callout(doc, "RÈGLE", "Pas de rubrique sans source, responsable et date. Pas de contenu pour remplir le template.")

    heading(doc, "Éditorial minimum viable", 2)
    table(doc, ["Format", "Cadence recommandée", "Responsable", "Validateur"], [
        ["Catalogue du moment", "À chaque nouvelle édition", "", ""],
        ["Sélection saisonnière", "1 fois / mois", "", ""],
        ["Conseil métier utile", "1 fois / mois", "", ""],
    ], [2700, 2400, 2130, 2130], size=8.8)

    heading(doc, "Contenus indispensables avant la maquette", 2)
    table(doc, ["À fournir", "Responsable", "Date", "Fait"], [
        ["Liste finale des familles produits", "", "", "[ ]"],
        ["Services et zone de livraison", "", "", "[ ]"],
        ["4 chiffres / preuves vérifiés", "", "", "[ ]"],
        ["Catalogues + couvertures + dates", "", "", "[ ]"],
        ["Photos réelles : équipe, livraison, produits, clients", "", "", "[ ]"],
        ["Coordonnées, horaires et destinataire des formulaires", "", "", "[ ]"],
    ], [4700, 2200, 1500, 960], size=8.6)

    heading(doc, "Décisions finales", 2)
    table(doc, ["Sujet", "Décision en une phrase", "Responsable"], [
        ["Cible", "", ""],
        ["Promesse", "", ""],
        ["CTA", "Devenir client", ""],
        ["Accueil", "8 blocs validés / corrections consignées", ""],
        ["Conversion", "Formulaire et délai de rappel validés", ""],
        ["Fiche produit", "Structure, données et CTA validés", ""],
    ], [1900, 5200, 2260], size=8.6)

    callout(doc, "PROCHAINE ÉTAPE", "Sous 48 h : consolider les réponses. Ensuite : wireframes de l’accueil, de la fiche produit et du parcours Devenir client + backlog des contenus.", fill=SAGE)

    props = doc.core_properties
    props.title = "Atelier UX express — Cash Alimentaire"
    props.subject = "Support de décision et de mise en application rapide"
    props.author = "Cash Alimentaire"
    props.comments = "Version express issue du benchmark distribution alimentaire B2B."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
